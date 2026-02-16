from __future__ import annotations

import io
from pathlib import Path
from typing import Iterable

from docx import Document as DocxDocument
from docx.text.run import Run


class WordParser:
    def load_docx(self, file_bytes: io.BytesIO) -> DocxDocument:
        file_bytes.seek(0)
        return DocxDocument(file_bytes)

    def load_with_conversion(
        self,
        file_bytes: io.BytesIO,
        extension: str,
        converter,
        tmp_dir: Path,
    ) -> DocxDocument:
        if extension in {"doc", "docm"}:
            file_bytes = converter.convert_to_docx(file_bytes, extension, tmp_dir)
        return self.load_docx(file_bytes)

    def iter_runs(self, doc: DocxDocument) -> list[Run]:
        runs: list[Run] = []
        for paragraph in doc.paragraphs:
            runs.extend([run for run in paragraph.runs if run.text])

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        runs.extend([run for run in paragraph.runs if run.text])

        for section in doc.sections:
            runs.extend(self._runs_from_paragraphs(section.header.paragraphs))
            runs.extend(self._runs_from_paragraphs(section.footer.paragraphs))

        return runs

    def apply_translations(self, runs: Iterable[Run], translations: Iterable[str]) -> None:
        for run, translated in zip(runs, translations, strict=False):
            run.text = translated

    def to_bytes(self, doc: DocxDocument) -> io.BytesIO:
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    @staticmethod
    def _runs_from_paragraphs(paragraphs) -> list[Run]:
        runs: list[Run] = []
        for paragraph in paragraphs:
            runs.extend([run for run in paragraph.runs if run.text])
        return runs
